"""
Markdown to DOCX Converter
Pure Python implementation to convert markdown to Word documents

© 2025-2030 Ashutosh Sinha, ajsinha@gmail.com, https://www.github.com/ajsinha/abhikarta

Required packages:
    pip install python-docx markdown
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
import markdown
from html.parser import HTMLParser


class MarkdownToDocxConverter:
    """Convert Markdown to DOCX with formatting"""

    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Setup custom styles for the document"""
        styles = self.doc.styles

        # Heading 1 style
        if 'Heading 1' not in [s.name for s in styles]:
            h1_style = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            h1_style.font.size = Pt(16)
            h1_style.font.bold = True
            h1_style.font.color.rgb = RGBColor(13, 110, 253)  # Blue
        else:
            h1_style = styles['Heading 1']
            h1_style.font.size = Pt(16)
            h1_style.font.bold = True
            h1_style.font.color.rgb = RGBColor(13, 110, 253)

        # Heading 2 style
        if 'Heading 2' not in [s.name for s in styles]:
            h2_style = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
            h2_style.font.size = Pt(14)
            h2_style.font.bold = True
            h2_style.font.color.rgb = RGBColor(73, 80, 87)  # Dark gray
        else:
            h2_style = styles['Heading 2']
            h2_style.font.size = Pt(14)
            h2_style.font.bold = True
            h2_style.font.color.rgb = RGBColor(73, 80, 87)

        # Code style
        if 'Code' not in [s.name for s in styles]:
            code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
            code_style.font.name = 'Courier New'
            code_style.font.size = Pt(10)

    def convert(self, markdown_text):
        """
        Convert markdown text to a docx Document

        Args:
            markdown_text (str): Markdown formatted text

        Returns:
            Document: python-docx Document object
        """
        lines = markdown_text.split('\n')
        i = 0
        in_code_block = False
        code_lines = []
        in_list = False
        list_items = []

        while i < len(lines):
            line = lines[i]

            # Handle code blocks
            if line.strip().startswith('```'):
                if in_code_block:
                    # End of code block
                    code_text = '\n'.join(code_lines)
                    p = self.doc.add_paragraph(code_text, style='Code')
                    p.paragraph_format.left_indent = Inches(0.5)
                    code_lines = []
                    in_code_block = False
                else:
                    # Start of code block
                    in_code_block = True
                i += 1
                continue

            if in_code_block:
                code_lines.append(line)
                i += 1
                continue

            # Handle tables
            if line.strip().startswith('|') and '|' in line[1:]:
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    # Skip separator lines
                    if not all(c in '|-: ' for c in lines[i].strip()):
                        cells = [cell.strip() for cell in lines[i].split('|')[1:-1]]
                        if cells:
                            table_lines.append(cells)
                    i += 1

                if table_lines:
                    self._add_table(table_lines)
                continue

            # Handle unordered lists
            bullet_match = re.match(r'^[\s]*[-*]\s+(.+)$', line)
            if bullet_match:
                if not in_list:
                    in_list = True
                    list_items = []
                list_items.append(bullet_match.group(1))
                i += 1
                continue

            # Handle ordered lists
            numbered_match = re.match(r'^[\s]*\d+\.\s+(.+)$', line)
            if numbered_match:
                if not in_list:
                    in_list = True
                    list_items = []
                list_items.append(numbered_match.group(1))
                i += 1
                continue

            # End of list
            if in_list and line.strip() == '':
                for item in list_items:
                    self._add_paragraph_with_formatting(item, is_list=True)
                in_list = False
                list_items = []
                i += 1
                continue

            # Handle headings
            if line.startswith('# '):
                self.doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith('## '):
                self.doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith('### '):
                self.doc.add_heading(line[4:].strip(), level=3)
            elif line.startswith('#### '):
                self.doc.add_heading(line[5:].strip(), level=4)

            # Handle blockquotes
            elif line.strip().startswith('> '):
                quote_text = line.strip()[2:]
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(12)
                run = p.add_run(quote_text)
                run.italic = True
                run.font.color.rgb = RGBColor(108, 117, 125)

            # Handle horizontal rules
            elif line.strip() in ['---', '***', '___']:
                self.doc.add_paragraph('_' * 50)

            # Handle empty lines
            elif line.strip() == '':
                self.doc.add_paragraph()

            # Handle regular paragraphs
            elif line.strip():
                self._add_paragraph_with_formatting(line)

            i += 1

        # Add any remaining list items
        if in_list and list_items:
            for item in list_items:
                self._add_paragraph_with_formatting(item, is_list=True)

        return self.doc

    def _add_table(self, table_lines):
        """Add a table to the document"""
        if not table_lines:
            return

        rows = len(table_lines)
        cols = len(table_lines[0])

        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Light Grid Accent 1'

        for i, row_data in enumerate(table_lines):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                cell = row.cells[j]
                # Parse inline formatting for cell text
                self._add_formatted_text(cell.paragraphs[0], cell_text)

                # Make header row bold
                if i == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True

    def _add_paragraph_with_formatting(self, text, is_list=False):
        """Add a paragraph with inline formatting (bold, italic, code, links)"""
        p = self.doc.add_paragraph()

        if is_list:
            p.style = 'List Bullet'

        self._add_formatted_text(p, text)

    def _add_formatted_text(self, paragraph, text):
        """Parse and add formatted text to a paragraph"""
        # Pattern to match bold, italic, code, and links
        pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[([^\]]+)\]\(([^)]+)\))'

        last_end = 0

        for match in re.finditer(pattern, text):
            # Add text before the match
            if match.start() > last_end:
                paragraph.add_run(text[last_end:match.start()])

            matched_text = match.group(0)

            # Bold
            if matched_text.startswith('**') and matched_text.endswith('**'):
                run = paragraph.add_run(matched_text[2:-2])
                run.bold = True

            # Italic
            elif matched_text.startswith('*') and matched_text.endswith('*'):
                run = paragraph.add_run(matched_text[1:-1])
                run.italic = True

            # Inline code
            elif matched_text.startswith('`') and matched_text.endswith('`'):
                run = paragraph.add_run(matched_text[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(10)

            # Links
            elif matched_text.startswith('['):
                link_text = match.group(2)
                link_url = match.group(3)
                run = paragraph.add_run(link_text)
                run.font.color.rgb = RGBColor(0, 102, 204)
                run.underline = True

            last_end = match.end()

        # Add remaining text
        if last_end < len(text):
            paragraph.add_run(text[last_end:])

    def save(self, filename):
        """Save the document to a file"""
        self.doc.save(filename)


def convert_markdown_to_docx(markdown_text, output_filename):
    """
    Convert markdown text to a DOCX file

    Args:
        markdown_text (str): Markdown formatted text
        output_filename (str): Path to save the DOCX file
    """
    converter = MarkdownToDocxConverter()
    converter.convert(markdown_text)
    converter.save(output_filename)
    return output_filename


# CLI usage
if __name__ == '__main__':
    import sys

    if len(sys.argv) < 3:
        print("Usage: python markdown_to_docx.py <input.md> <output.docx>")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]

    with open(input_file, 'r', encoding='utf-8') as f:
        markdown_text = f.read()

    convert_markdown_to_docx(markdown_text, output_file)
    print(f"Converted {input_file} to {output_file}")