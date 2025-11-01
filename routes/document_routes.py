"""
Document Routes
Document generation and management routes

© 2025-2030 Ashutosh Sinha, ajsinha@gmail.com, https://www.github.com/ajsinha/abhikarta
"""

import os
import re
from flask import session, render_template, request, jsonify
from werkzeug.utils import secure_filename

ALLOWED_EXTENSIONS = {'txt', 'pdf', 'doc', 'docx', 'md', 'csv', 'json'}
MAX_FILES = 10


class DocumentRoutes:
    """Document generation routes handler"""

    def __init__(self, app, user_registry, login_required):
        """
        Initialize document routes

        Args:
            app: Flask application instance
            user_registry: User registry instance
            login_required: Login required decorator
        """
        self.app = app
        self.user_registry = user_registry
        self.login_required = login_required
        self.register_routes()

    def register_routes(self):
        """Register all document routes"""

        @self.app.route('/document-generation')
        @self.login_required
        def document_generation():
            """Document generation page"""
            user = self.user_registry.get_user(session['user_id'])
            return render_template('document_generate.html', user=user)

        @self.app.route('/api/document/list-sessions', methods=['GET'])
        @self.login_required
        def list_document_sessions():
            """List all document generation sessions for the current user"""
            user = self.user_registry.get_user(session['user_id'])

            # Path to user's docgen folder
            docgen_path = os.path.join('data', 'uploads', str(user.user_id), 'docgen')

            try:
                sessions = []

                # Check if docgen folder exists
                if os.path.exists(docgen_path):
                    # Get all directories in docgen folder
                    for item in os.listdir(docgen_path):
                        item_path = os.path.join(docgen_path, item)
                        if os.path.isdir(item_path):
                            # Get file count in inbox
                            inbox_path = os.path.join(item_path, 'inbox')
                            file_count = 0
                            if os.path.exists(inbox_path):
                                file_count = len([f for f in os.listdir(inbox_path)
                                                if os.path.isfile(os.path.join(inbox_path, f))
                                                and not f.startswith('.')])

                            # Get last modified time
                            modified_time = os.path.getmtime(item_path)

                            sessions.append({
                                'name': item,
                                'file_count': file_count,
                                'modified_time': modified_time
                            })

                    # Sort by modified time (most recent first)
                    sessions.sort(key=lambda x: x['modified_time'], reverse=True)

                return jsonify({
                    'success': True,
                    'sessions': sessions,
                    'count': len(sessions)
                })

            except Exception as e:
                return jsonify({'success': False, 'error': f'Failed to list sessions: {str(e)}'}), 500

        @self.app.route('/api/document/delete-session', methods=['POST'])
        @self.login_required
        def delete_document_session():
            """Delete a document generation session"""
            import shutil

            user = self.user_registry.get_user(session['user_id'])
            data = request.get_json()

            session_name = data.get('session_name', '').strip()

            if not session_name:
                return jsonify({'success': False, 'error': 'Session name is required'}), 400

            # Path to session folder
            session_path = os.path.join('data', 'uploads', str(user.user_id), 'docgen', session_name)

            try:
                # Check if session exists
                if not os.path.exists(session_path):
                    return jsonify({'success': False, 'error': 'Session not found'}), 404

                # Check if it's a directory (safety check)
                if not os.path.isdir(session_path):
                    return jsonify({'success': False, 'error': 'Invalid session path'}), 400

                # Delete the entire session folder
                shutil.rmtree(session_path)

                return jsonify({
                    'success': True,
                    'message': f'Session "{session_name}" deleted successfully'
                })

            except Exception as e:
                return jsonify({'success': False, 'error': f'Failed to delete session: {str(e)}'}), 500

        @self.app.route('/api/document/create-session', methods=['POST'])
        @self.login_required
        def create_document_session():
            """Create a new document generation session"""
            user = self.user_registry.get_user(session['user_id'])
            data = request.get_json()
            session_name = data.get('session_name', '').strip()

            # Validate session name
            if not session_name:
                return jsonify({'success': False, 'error': 'Session name is required'}), 400

            # Check if session name contains at least one alphabetic character
            if not re.search(r'[a-zA-Z]', session_name):
                return jsonify({'success': False, 'error': 'Session name must contain at least one alphabetic character'}), 400

            # Check if session name contains only alphanumerics and underscores
            if not re.match(r'^[a-zA-Z0-9_]+$', session_name):
                return jsonify({'success': False, 'error': 'Session name can only contain alphanumerics and underscores'}), 400

            # Create session folder structure
            base_path = os.path.join('data', 'uploads', str(user.user_id), 'docgen', session_name)
            inbox_path = os.path.join(base_path, 'inbox')
            outbox_path = os.path.join(base_path, 'outbox')
            servermessages_path = os.path.join(base_path, 'servermessages')

            try:
                # Check if session already exists
                session_exists = os.path.exists(base_path)

                # Create all three folders (will not overwrite existing folders)
                os.makedirs(inbox_path, exist_ok=True)
                os.makedirs(outbox_path, exist_ok=True)
                os.makedirs(servermessages_path, exist_ok=True)

                # Get list of existing files in inbox
                existing_files = []
                if os.path.exists(inbox_path):
                    existing_files = [f for f in os.listdir(inbox_path)
                                    if os.path.isfile(os.path.join(inbox_path, f))
                                    and not f.startswith('.')]

                message = f'Session "{session_name}" '
                if session_exists and existing_files:
                    message += f'loaded with {len(existing_files)} existing file(s)'
                elif session_exists:
                    message += 'loaded successfully'
                else:
                    message += 'created successfully'

                return jsonify({
                    'success': True,
                    'message': message,
                    'session_path': base_path,
                    'inbox_path': inbox_path,
                    'outbox_path': outbox_path,
                    'servermessages_path': servermessages_path,
                    'session_exists': session_exists,
                    'existing_files': existing_files
                })
            except Exception as e:
                return jsonify({'success': False, 'error': f'Failed to create session: {str(e)}'}), 500

        @self.app.route('/api/document/upload', methods=['POST'])
        @self.login_required
        def upload_documents():
            """Upload documents to session inbox folder"""
            user = self.user_registry.get_user(session['user_id'])
            session_name = request.form.get('session_name')

            if not session_name:
                return jsonify({'success': False, 'error': 'Session name is required'}), 400

            # Check if files were uploaded
            if 'files' not in request.files:
                return jsonify({'success': False, 'error': 'No files uploaded'}), 400

            files = request.files.getlist('files')

            # Check file count
            if len(files) > MAX_FILES:
                return jsonify({'success': False, 'error': f'Maximum {MAX_FILES} files allowed'}), 400

            # Path to inbox folder
            inbox_path = os.path.join('data', 'uploads', str(user.user_id), 'docgen', session_name, 'inbox')

            if not os.path.exists(inbox_path):
                return jsonify({'success': False, 'error': 'Session folder does not exist. Please create a session first.'}), 400

            uploaded_files = []

            try:
                for file in files:
                    if file.filename == '':
                        continue

                    if file and self._allowed_file(file.filename):
                        filename = secure_filename(file.filename)
                        file_path = os.path.join(inbox_path, filename)
                        file.save(file_path)
                        uploaded_files.append(filename)
                    else:
                        return jsonify({'success': False, 'error': f'File type not allowed: {file.filename}'}), 400

                return jsonify({
                    'success': True,
                    'message': f'{len(uploaded_files)} file(s) uploaded successfully',
                    'files': uploaded_files
                })
            except Exception as e:
                return jsonify({'success': False, 'error': f'Upload failed: {str(e)}'}), 500

        @self.app.route('/api/document/generate', methods=['POST'])
        @self.login_required
        def generate_document():
            """Generate document based on uploaded files and instructions"""
            user = self.user_registry.get_user(session['user_id'])
            data = request.get_json()

            session_name = data.get('session_name', '').strip()
            template = data.get('template', 'Default')
            instructions = data.get('instructions', '').strip()

            if not session_name:
                return jsonify({'success': False, 'error': 'Session name is required'}), 400

            # Get uploaded files from inbox
            inbox_path = os.path.join('data', 'uploads', str(user.user_id), 'docgen', session_name, 'inbox')

            if not os.path.exists(inbox_path):
                return jsonify({'success': False, 'error': 'Session not found'}), 404

            uploaded_files = [f for f in os.listdir(inbox_path) if os.path.isfile(os.path.join(inbox_path, f))]

            # Generate document content (hardcoded for now)
            generated_content = self._generate_document_content(session_name, template, instructions, uploaded_files)

            return jsonify({
                'success': True,
                'content': generated_content,
                'session_name': session_name,
                'template': template,
                'files_processed': len(uploaded_files)
            })

        @self.app.route('/api/document/delete-file', methods=['POST'])
        @self.login_required
        def delete_file():
            """Delete a file from the session inbox"""
            user = self.user_registry.get_user(session['user_id'])
            data = request.get_json()

            session_name = data.get('session_name', '').strip()
            filename = data.get('filename', '').strip()

            if not session_name:
                return jsonify({'success': False, 'error': 'Session name is required'}), 400

            if not filename:
                return jsonify({'success': False, 'error': 'Filename is required'}), 400

            # Path to the file in inbox
            inbox_path = os.path.join('data', 'uploads', str(user.user_id), 'docgen', session_name, 'inbox')
            file_path = os.path.join(inbox_path, secure_filename(filename))

            if not os.path.exists(inbox_path):
                return jsonify({'success': False, 'error': 'Session not found'}), 404

            if not os.path.exists(file_path):
                return jsonify({'success': False, 'error': 'File not found'}), 404

            try:
                # Delete the file
                os.remove(file_path)
                return jsonify({
                    'success': True,
                    'message': f'File "{filename}" deleted successfully'
                })
            except Exception as e:
                return jsonify({'success': False, 'error': f'Failed to delete file: {str(e)}'}), 500

        @self.app.route('/api/document/rerun', methods=['POST'])
        @self.login_required
        def rerun_generation():
            """Rerun document generation with current editor content"""
            user = self.user_registry.get_user(session['user_id'])
            data = request.get_json()

            session_name = data.get('session_name', '').strip()
            template = data.get('template', 'Default')
            instructions = data.get('instructions', '').strip()
            current_content = data.get('current_content', '').strip()

            if not session_name:
                return jsonify({'success': False, 'error': 'Session name is required'}), 400

            # Get uploaded files from inbox (for context)
            inbox_path = os.path.join('data', 'uploads', str(user.user_id), 'docgen', session_name, 'inbox')

            if not os.path.exists(inbox_path):
                return jsonify({'success': False, 'error': 'Session not found'}), 404

            uploaded_files = [f for f in os.listdir(inbox_path) if os.path.isfile(os.path.join(inbox_path, f))]

            # For now, return the same hardcoded content (can be enhanced later to process current_content)
            # In a real implementation, you might:
            # - Parse current_content and enhance it
            # - Use LLM to regenerate based on current_content
            # - Apply additional processing

            generated_content = self._generate_document_content(session_name, template, instructions, uploaded_files)

            return jsonify({
                'success': True,
                'content': generated_content,
                'session_name': session_name,
                'template': template,
                'message': 'Document regenerated (current implementation uses same template)'
            })

        @self.app.route('/api/document/download-word', methods=['POST'])
        @self.login_required
        def download_word():
            """Convert markdown content to Word document and return as download"""
            from flask import send_file
            import tempfile
            from io import BytesIO

            user = self.user_registry.get_user(session['user_id'])
            data = request.get_json()

            markdown_content = data.get('content', '').strip()
            filename = data.get('filename', 'document')

            if not markdown_content:
                return jsonify({'success': False, 'error': 'No content provided'}), 400

            try:
                # Try to use pypandoc if available (best quality conversion)
                try:
                    import pypandoc

                    # Create temporary markdown file
                    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as md_file:
                        md_file.write(markdown_content)
                        md_path = md_file.name

                    # Create temporary output file
                    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as docx_file:
                        docx_path = docx_file.name

                    # Convert using pandoc
                    pypandoc.convert_file(md_path, 'docx', outputfile=docx_path)

                    # Clean up markdown file
                    os.unlink(md_path)

                    # Return the Word document
                    return send_file(
                        docx_path,
                        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        as_attachment=True,
                        download_name=f'{filename}.docx'
                    )

                except (ImportError, RuntimeError):
                    # Fallback to python-docx with markdown parsing
                    from docx import Document
                    from docx.shared import Pt, Inches
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    import re

                    doc = Document()

                    # Process markdown line by line
                    lines = markdown_content.split('\n')
                    i = 0
                    while i < len(lines):
                        line = lines[i].rstrip()

                        if not line:
                            i += 1
                            continue

                        # Headers
                        if line.startswith('# '):
                            doc.add_heading(line[2:], level=1)
                        elif line.startswith('## '):
                            doc.add_heading(line[3:], level=2)
                        elif line.startswith('### '):
                            doc.add_heading(line[4:], level=3)
                        elif line.startswith('#### '):
                            doc.add_heading(line[5:], level=4)

                        # Bullet lists
                        elif line.startswith('- ') or line.startswith('* '):
                            doc.add_paragraph(line[2:], style='List Bullet')

                        # Numbered lists
                        elif re.match(r'^\d+\.\s', line):
                            text = re.sub(r'^\d+\.\s', '', line)
                            doc.add_paragraph(text, style='List Number')

                        # Tables (basic support)
                        elif line.startswith('|') and i + 1 < len(lines) and lines[i + 1].startswith('|'):
                            # Parse table
                            table_lines = []
                            while i < len(lines) and lines[i].startswith('|'):
                                table_lines.append(lines[i])
                                i += 1

                            if len(table_lines) > 2:  # At least header + separator + one row
                                # Extract cells
                                rows = []
                                for tline in table_lines:
                                    cells = [cell.strip() for cell in tline.split('|')[1:-1]]
                                    # Skip separator line
                                    if not all(c.replace('-', '').replace(':', '').strip() == '' for c in cells):
                                        rows.append(cells)

                                if rows:
                                    # Create table
                                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                                    table.style = 'Light Grid Accent 1'

                                    for row_idx, row_cells in enumerate(rows):
                                        for col_idx, cell_text in enumerate(row_cells):
                                            table.rows[row_idx].cells[col_idx].text = cell_text

                            continue  # Already incremented i

                        # Horizontal rules
                        elif line.startswith('---') or line.startswith('***'):
                            doc.add_paragraph()
                            p = doc.add_paragraph()
                            p.add_run('_' * 50)
                            doc.add_paragraph()

                        # Code blocks
                        elif line.startswith('```'):
                            i += 1
                            code_lines = []
                            while i < len(lines) and not lines[i].startswith('```'):
                                code_lines.append(lines[i])
                                i += 1

                            if code_lines:
                                p = doc.add_paragraph('\n'.join(code_lines))
                                p.style = 'No Spacing'
                                run = p.runs[0]
                                run.font.name = 'Courier New'
                                run.font.size = Pt(9)

                        # Bold text
                        elif '**' in line:
                            p = doc.add_paragraph()
                            parts = re.split(r'(\*\*[^*]+\*\*)', line)
                            for part in parts:
                                if part.startswith('**') and part.endswith('**'):
                                    run = p.add_run(part[2:-2])
                                    run.bold = True
                                else:
                                    p.add_run(part)

                        # Regular paragraph
                        else:
                            doc.add_paragraph(line)

                        i += 1

                    # Save to BytesIO
                    docx_io = BytesIO()
                    doc.save(docx_io)
                    docx_io.seek(0)

                    return send_file(
                        docx_io,
                        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        as_attachment=True,
                        download_name=f'{filename}.docx'
                    )

            except Exception as e:
                return jsonify({'success': False, 'error': f'Conversion failed: {str(e)}'}), 500

    def _allowed_file(self, filename):
        """Check if file extension is allowed"""
        return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

    def _generate_document_content(self, session_name, template, instructions, files):
        """Generate hardcoded document content"""

        def load_sample_content():
            with open("data/sample/sample_pfe_model_documentation.md", "r") as file:
                content = file.read()
                return content

        if template == "CCR PFE":
            content = load_sample_content()
        else:  # Default template
            content = load_sample_content()

        return content